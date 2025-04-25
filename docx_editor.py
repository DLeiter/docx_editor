#!/usr/bin/env python3
"""
DOCX Editor - A program to edit Word documents with advanced features
"""
import os
import sys
import io
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, colorchooser, font
from tkinter.constants import *
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table, _Cell
from PIL import Image, ImageTk
import base64
from io import BytesIO

# Tooltip class implementation for showing tooltips on hover
class ToolTip(object):
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip_window = None
        self.widget.bind("<Enter>", self.enter)
        self.widget.bind("<Leave>", self.leave)
        self.widget.bind("<ButtonPress>", self.leave)
    
    def enter(self, event=None):
        x, y, _, _ = self.widget.bbox("insert")
        x += self.widget.winfo_rootx() + 25
        y += self.widget.winfo_rooty() + 25
        
        # Creates a toplevel window
        self.tooltip_window = tk.Toplevel(self.widget)
        # Leaves only the label and removes the app window
        self.tooltip_window.wm_overrideredirect(True)
        self.tooltip_window.wm_geometry(f"+{x}+{y}")
        
        # Creates the label for the tooltip
        label = tk.Label(self.tooltip_window, text=self.text, background="#FFFFCC",
                     relief="solid", borderwidth=1, font=("Arial", "9", "normal"))
        label.pack(padx=2, pady=2)
    
    def leave(self, event=None):
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None

class DocxEditor:
    def __init__(self, root):
        self.root = root
        self.root.title("DOCX Editor")
        self.root.geometry("1000x800")
        
        # Initialize variables
        self.document = None
        self.current_file = ""
        self.tables = []
        self.document_images = []
        
        # Text formatting state variables
        self.current_font_family = "Arial"
        self.current_font_size = 12
        self.current_bold = False
        self.current_italic = False
        self.current_underline = False
        self.current_color = "#000000"
        self.current_highlight = "#FFFFFF"
        self.current_alignment = "left"
        self.current_style = "Normal"#FFFF00"  # Yellow
        
        # Table management
        self.tables = []
        self.current_table = None
        
        # Paragraph styling
        self.paragraph_styles = []
        
        # Document sections and headers/footers tracking
        self.sections = []
        self.headers = {}
        self.footers = {}
        
        # Tooltip dictionary to keep track of tooltips
        self.tooltips = {}
        
        # Create all widgets and interface elements
        self.create_menu()
        self.create_widgets()
        self.create_formatting_toolbar()
        self.create_tabs()
        
        # Set up keyboard shortcuts
        self._setup_keyboard_shortcuts()
        
        # Set initial status
        self.status_var.set("Ready")
        
        # Auto-load the first example document if available
        self.auto_load_example_document()
        
    def create_menu(self):
        menubar = tk.Menu(self.root)
        
        # File menu
        filemenu = tk.Menu(menubar, tearoff=0)
        filemenu.add_command(label="New", command=self.new_document)
        filemenu.add_command(label="Open", command=self.open_file)
        filemenu.add_command(label="Save", command=self.save_file)
        filemenu.add_command(label="Save As", command=self.save_file_as)
        filemenu.add_separator()
        filemenu.add_command(label="Export as PDF", command=self.export_pdf)
        filemenu.add_separator()
        filemenu.add_command(label="Exit", command=self.root.quit)
        menubar.add_cascade(label="File", menu=filemenu)
        
        # Edit menu
        editmenu = tk.Menu(menubar, tearoff=0)
        editmenu.add_command(label="Undo", command=self.undo)
        editmenu.add_command(label="Redo", command=self.redo)
        editmenu.add_separator()
        editmenu.add_command(label="Cut", command=lambda: self.text_editor.event_generate("<<Cut>>"))
        editmenu.add_command(label="Copy", command=lambda: self.text_editor.event_generate("<<Copy>>"))
        editmenu.add_command(label="Paste", command=lambda: self.text_editor.event_generate("<<Paste>>"))
        editmenu.add_separator()
        editmenu.add_command(label="Find/Replace", command=self.find_replace_dialog)
        editmenu.add_separator()
        editmenu.add_command(label="Clear All", command=self.clear_text)
        menubar.add_cascade(label="Edit", menu=editmenu)
        
        # Format menu
        formatmenu = tk.Menu(menubar, tearoff=0)
        
        # Text formatting submenu
        text_formatting_menu = tk.Menu(formatmenu, tearoff=0)
        text_formatting_menu.add_command(label="Font...", command=self.font_dialog)
        text_formatting_menu.add_command(label="Text Color...", command=self.text_color_dialog)
        text_formatting_menu.add_separator()
        text_formatting_menu.add_command(label="Bold", command=self.toggle_bold)
        text_formatting_menu.add_command(label="Italic", command=self.toggle_italic)
        text_formatting_menu.add_command(label="Underline", command=self.toggle_underline)
        formatmenu.add_cascade(label="Text", menu=text_formatting_menu)
        
        # Paragraph formatting submenu
        paragraph_menu = tk.Menu(formatmenu, tearoff=0)
        paragraph_menu.add_command(label="Align Left", command=lambda: self.set_alignment("left"))
        paragraph_menu.add_command(label="Align Center", command=lambda: self.set_alignment("center"))
        paragraph_menu.add_command(label="Align Right", command=lambda: self.set_alignment("right"))
        paragraph_menu.add_command(label="Justify", command=lambda: self.set_alignment("justify"))
        paragraph_menu.add_separator()
        paragraph_menu.add_command(label="Paragraph Styles...", command=self.paragraph_style_dialog)
        formatmenu.add_cascade(label="Paragraph", menu=paragraph_menu)
        
        menubar.add_cascade(label="Format", menu=formatmenu)
        
        # Insert menu
        insertmenu = tk.Menu(menubar, tearoff=0)
        insertmenu.add_command(label="Image...", command=self.insert_image)
        insertmenu.add_command(label="Table...", command=self.insert_table_dialog)
        insertmenu.add_separator()
        insertmenu.add_command(label="Page Break", command=self.insert_page_break)
        insertmenu.add_command(label="Section Break", command=self.insert_section_break)
        insertmenu.add_separator()
        insertmenu.add_command(label="Hyperlink...", command=self.insert_hyperlink)
        insertmenu.add_command(label="Table of Contents", command=self.insert_toc)
        insertmenu.add_separator()
        insertmenu.add_command(label="Header", command=self.edit_header)
        insertmenu.add_command(label="Footer", command=self.edit_footer)
        menubar.add_cascade(label="Insert", menu=insertmenu)
        
        # Table menu
        tablemenu = tk.Menu(menubar, tearoff=0)
        tablemenu.add_command(label="Insert Table...", command=self.insert_table_dialog)
        tablemenu.add_command(label="Edit Table...", command=self.edit_table_dialog)
        tablemenu.add_separator()
        tablemenu.add_command(label="Add Row", command=self.add_table_row)
        tablemenu.add_command(label="Add Column", command=self.add_table_column)
        tablemenu.add_separator()
        tablemenu.add_command(label="Delete Row", command=self.delete_table_row)
        tablemenu.add_command(label="Delete Column", command=self.delete_table_column)
        menubar.add_cascade(label="Table", menu=tablemenu)
        
        # View menu
        viewmenu = tk.Menu(menubar, tearoff=0)
        viewmenu.add_command(label="Document Properties", command=self.document_properties)
        viewmenu.add_separator()
        viewmenu.add_command(label="Zoom In", command=self.zoom_in)
        viewmenu.add_command(label="Zoom Out", command=self.zoom_out)
        viewmenu.add_command(label="Reset Zoom", command=self.reset_zoom)
        menubar.add_cascade(label="View", menu=viewmenu)
        
        # Help menu
        helpmenu = tk.Menu(menubar, tearoff=0)
        helpmenu.add_command(label="About", command=self.show_about)
        helpmenu.add_command(label="Help", command=self.show_help)
        menubar.add_cascade(label="Help", menu=helpmenu)
        
        self.root.config(menu=menubar)
    
    def create_formatting_toolbar(self):
        # Create a frame for the formatting toolbar
        self.format_frame = ttk.Frame(self.root)
        self.format_frame.pack(fill=tk.X, padx=5, pady=2)
        
        # Font family selector
        ttk.Label(self.format_frame, text="Font:").pack(side=tk.LEFT, padx=2)
        self.font_family = ttk.Combobox(self.format_frame, width=15, state="readonly")
        self.font_family["values"] = sorted(font.families())
        self.font_family.current(self.font_family["values"].index(self.current_font_family) if self.current_font_family in self.font_family["values"] else 0)
        self.font_family.bind("<<ComboboxSelected>>", self.change_font_family)
        self.font_family.pack(side=tk.LEFT, padx=2)
        
        # Font size selector
        ttk.Label(self.format_frame, text="Size:").pack(side=tk.LEFT, padx=2)
        self.font_size = ttk.Combobox(self.format_frame, width=5, state="readonly")
        self.font_size["values"] = [8, 9, 10, 11, 12, 14, 16, 18, 20, 22, 24, 26, 28, 36, 48, 72]
        self.font_size.current(self.font_size["values"].index(self.current_font_size) if self.current_font_size in self.font_size["values"] else 3)  # Default to 11
        self.font_size.bind("<<ComboboxSelected>>", self.change_font_size)
        self.font_size.pack(side=tk.LEFT, padx=2)
        
        # Style dropdown
        ttk.Label(self.format_frame, text="Style:").pack(side=tk.LEFT, padx=2)
        self.style_combo = ttk.Combobox(self.format_frame, width=15, state="readonly")
        self.style_combo["values"] = ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle"]
        self.style_combo.current(0)
        self.style_combo.bind("<<ComboboxSelected>>", self.apply_style)
        self.style_combo.pack(side=tk.LEFT, padx=2)
        
        # Separator
        ttk.Separator(self.format_frame, orient=tk.VERTICAL).pack(side=tk.LEFT, padx=5, fill=tk.Y)
        
        # Bold button
        self.bold_icon = tk.PhotoImage(data=self._get_bold_icon())
        self.bold_button = ttk.Button(self.format_frame, image=self.bold_icon, width=3, command=self.toggle_bold)
        self._create_tooltip(self.bold_button, "Bold (Ctrl+B)")
        self.bold_button.pack(side=tk.LEFT, padx=2)
        
        # Italic button
        self.italic_icon = tk.PhotoImage(data=self._get_italic_icon())
        self.italic_button = ttk.Button(self.format_frame, image=self.italic_icon, width=3, command=self.toggle_italic)
        self._create_tooltip(self.italic_button, "Italic (Ctrl+I)")
        self.italic_button.pack(side=tk.LEFT, padx=2)
        
        # Underline button
        self.underline_icon = tk.PhotoImage(data=self._get_underline_icon())
        self.underline_button = ttk.Button(self.format_frame, image=self.underline_icon, width=3, command=self.toggle_underline)
        self._create_tooltip(self.underline_button, "Underline (Ctrl+U)")
        self.underline_button.pack(side=tk.LEFT, padx=2)
        
        # Text color button - use standard tk.Button instead of ttk for direct color support
        self.color_button = tk.Button(self.format_frame, text="A", width=3, fg=self.current_color, command=self.text_color_dialog)
        self._create_tooltip(self.color_button, "Text Color")
        self.color_button.pack(side=tk.LEFT, padx=2)
        
        # Highlight color button - use standard tk.Button instead of ttk for direct color support
        self.highlight_button = tk.Button(self.format_frame, text="H", width=3, bg="yellow", command=self.highlight_color_dialog)
        self._create_tooltip(self.highlight_button, "Highlight Color")
        self.highlight_button.pack(side=tk.LEFT, padx=2)
        
        # Separator
        ttk.Separator(self.format_frame, orient=tk.VERTICAL).pack(side=tk.LEFT, padx=5, fill=tk.Y)
        
        # Alignment buttons
        self.align_left_icon = tk.PhotoImage(data=self._get_align_left_icon())
        self.align_left_button = ttk.Button(self.format_frame, image=self.align_left_icon, width=3, command=lambda: self.set_alignment("left"))
        self._create_tooltip(self.align_left_button, "Align Left (Ctrl+L)")
        self.align_left_button.pack(side=tk.LEFT, padx=2)
        
        self.align_center_icon = tk.PhotoImage(data=self._get_align_center_icon())
        self.align_center_button = ttk.Button(self.format_frame, image=self.align_center_icon, width=3, command=lambda: self.set_alignment("center"))
        self._create_tooltip(self.align_center_button, "Align Center (Ctrl+E)")
        self.align_center_button.pack(side=tk.LEFT, padx=2)
        
        self.align_right_icon = tk.PhotoImage(data=self._get_align_right_icon())
        self.align_right_button = ttk.Button(self.format_frame, image=self.align_right_icon, width=3, command=lambda: self.set_alignment("right"))
        self._create_tooltip(self.align_right_button, "Align Right (Ctrl+R)")
        self.align_right_button.pack(side=tk.LEFT, padx=2)
        
        # Separator
        ttk.Separator(self.format_frame, orient=tk.VERTICAL).pack(side=tk.LEFT, padx=5, fill=tk.Y)
        
        # Bullet list button
        self.bullet_list_button = ttk.Button(self.format_frame, text="•", width=3, command=self.insert_bullet_list)
        self._create_tooltip(self.bullet_list_button, "Bullet List")
        self.bullet_list_button.pack(side=tk.LEFT, padx=2)
        
        # Numbered list button
        self.numbered_list_button = ttk.Button(self.format_frame, text="1.", width=3, command=self.insert_numbered_list)
        self._create_tooltip(self.numbered_list_button, "Numbered List")
        self.numbered_list_button.pack(side=tk.LEFT, padx=2)
        
        # Separator
        ttk.Separator(self.format_frame, orient=tk.VERTICAL).pack(side=tk.LEFT, padx=5, fill=tk.Y)
        
        # Insert image button
        self.insert_image_button = ttk.Button(self.format_frame, text="Insert Image", command=self.insert_image)
        self._create_tooltip(self.insert_image_button, "Insert Image")
        self.insert_image_button.pack(side=tk.LEFT, padx=2)
        
        # Insert table button
        self.insert_table_button = ttk.Button(self.format_frame, text="Insert Table", command=self.insert_table_dialog)
        self._create_tooltip(self.insert_table_button, "Insert Table")
        self.insert_table_button.pack(side=tk.LEFT, padx=2)
        
        # Separator
        ttk.Separator(self.format_frame, orient=tk.VERTICAL).pack(side=tk.LEFT, padx=10, fill=tk.Y)
        
        # Heading navigation
        ttk.Label(self.format_frame, text="Go to:").pack(side=tk.LEFT, padx=5)
        self.heading_nav = ttk.Combobox(self.format_frame, width=25, state="readonly")
        self.heading_nav.pack(side=tk.LEFT, padx=2)
        self.heading_nav.bind("<<ComboboxSelected>>", self.navigate_to_heading)

    def create_tabs(self):
        # Create notebook (tabbed interface)
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(expand=True, fill=tk.BOTH, padx=5, pady=5)
        
        # Create document tab (main editor tab)
        self.document_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.document_frame, text="Document")
        
        # Create text editor in the document tab
        self.text_editor = scrolledtext.ScrolledText(self.document_frame, wrap=tk.WORD, undo=True)
        self.text_editor.pack(expand=True, fill=tk.BOTH)
        
        # Create document structure tab
        self.structure_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.structure_frame, text="Document Structure")
        
        # Create structure tree view
        self.structure_tree = ttk.Treeview(self.structure_frame)
        self.structure_tree.pack(expand=True, fill=tk.BOTH)
        self.structure_tree.heading("#0", text="Document Elements")
        
        # Create document properties tab
        self.properties_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.properties_frame, text="Properties")

    def create_widgets(self):
        # Frame for buttons
        self.button_frame = tk.Frame(self.root)
        self.button_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # New button
        self.new_button = ttk.Button(self.button_frame, text="New Document", command=self.new_document)
        self.new_button.pack(side=tk.LEFT, padx=5)
        
        # Open button
        self.open_button = ttk.Button(self.button_frame, text="Open Document", command=self.open_file)
        self.open_button.pack(side=tk.LEFT, padx=5)
        
        # Save button
        self.save_button = ttk.Button(self.button_frame, text="Save Document", command=self.save_file)
        self.save_button.pack(side=tk.LEFT, padx=5)
        
        # Export PDF button
        self.export_pdf_button = ttk.Button(self.button_frame, text="Export PDF", command=self.export_pdf)
        self.export_pdf_button.pack(side=tk.LEFT, padx=5)
        
        # Status bar
        self.status_var = tk.StringVar()
        self.status_var.set("No file open")
        self.status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
    
    # Document operation methods
    def new_document(self):
        # Check if there are unsaved changes
        result = True
        if self.document is not None:
            result = messagebox.askyesno("Unsaved Changes",
                                        "You have unsaved changes. Do you want to create a new document anyway?")
        
        if result:
            # Clear the editor and create a new document
            self.document = Document()
            self.current_file = None
            self.text_editor.delete(1.0, tk.END)
            self.document_images = []
            self.tables = []
            self.update_document_structure()
            self.status_var.set("New document created")
    
    def open_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if file_path:
            if self.load_document(file_path):
                self.status_var.set(f"Opened: {os.path.basename(file_path)}")
                messagebox.showinfo("Success", f"Opened {os.path.basename(file_path)}")
    
    def save_file(self):
        if not self.current_file:
            self.save_file_as()
        else:
            self.save_document(self.current_file)
    
    def save_file_as(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if file_path:
            self.save_document(file_path)
            self.current_file = file_path
    
    def save_document(self, file_path):
        try:
            # Create a new document if we don't have one
            if not self.document:
                self.document = Document()
            
            # Get text content from editor
            text_content = self.text_editor.get(1.0, tk.END)
            lines = text_content.split('\n')
            
            # Clear existing content
            for para in list(self.document.paragraphs):
                p = para._element
                p.getparent().remove(p)
                para._p = para._element = None
            
            # Track current section for handling breaks
            current_section = 0
            
            # Process each line and add to document
            skip_lines = []
            for i, line in enumerate(lines):
                # Skip table placeholder lines
                if line.startswith('[TABLE ') and line.endswith(']'):
                    continue
                    
                # Skip section break markers (we handle those separately)
                if line == "[SECTION BREAK]":
                    # Add a section break
                    self.document.add_section()
                    current_section += 1
                    continue
                    
                # Skip TOC placeholders (we handle those separately)
                if line.startswith('[TABLE OF CONTENTS:') and line.endswith(']'):
                    # Extract TOC info and add it
                    toc_title = line.split(':')[1].split(',')[0].strip()
                    # Add heading for TOC
                    self.document.add_heading(toc_title, level=1)
                    # Add blank paragraph that will be populated with TOC field
                    self.document.add_paragraph()
                    continue
                
                # Process normal paragraph
                paragraph = self.document.add_paragraph(line)
                
                # Detect heading styles by pattern
                if line.startswith('# '):
                    paragraph.style = 'Heading 1'
                    paragraph.text = line[2:]  # Remove the # marker
                elif line.startswith('## '):
                    paragraph.style = 'Heading 2'
                    paragraph.text = line[3:]  # Remove the ## marker
                elif line.startswith('### '):
                    paragraph.style = 'Heading 3'
                    paragraph.text = line[4:]  # Remove the ### marker
                
                # Apply paragraph formatting if default values are changed
                if self.current_alignment != "left":
                    if self.current_alignment == "center":
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif self.current_alignment == "right":
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    elif self.current_alignment == "justify":
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            # Re-add tables
            if hasattr(self, 'tables') and self.tables:
                for table in self.tables:
                    # Copy the table to the new document
                    tbl = table._tbl
                    self.document.add_paragraph()  # Add a paragraph to separate tables
                    new_p = self.document.add_paragraph()  # Paragraph to hold the table
                    new_p._p.addnext(tbl)
            
            # Add headers and footers if we have them
            if hasattr(self, 'document') and hasattr(self.document, 'sections'):
                for i, section in enumerate(self.document.sections):
                    # Add header if we have one for this section
                    if hasattr(self, 'headers') and i in self.headers:
                        header_content = self.headers[i]
                        # If it's a string, add it as paragraph text
                        if isinstance(header_content, str):
                            section.header.paragraphs[0].text = header_content
                    
                    # Add footer if we have one for this section
                    if hasattr(self, 'footers') and i in self.footers:
                        footer_content = self.footers[i]
                        # Handle different types of content
                        if isinstance(footer_content, dict):
                            # Dictionary with text and page numbers flag
                            section.footer.paragraphs[0].text = footer_content.get('text', '')
                            if footer_content.get('page_numbers', False):
                                # Add a paragraph with page number field
                                footer_para = section.footer.add_paragraph()
                                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                run = footer_para.add_run()
                                # Add page number field code
                                fldChar = OxmlElement('w:fldChar')
                                fldChar.set(qn('w:fldCharType'), 'begin')
                                run._element.append(fldChar)
                                
                                instrText = OxmlElement('w:instrText')
                                instrText.text = 'PAGE'
                                run._element.append(instrText)
                                
                                fldChar = OxmlElement('w:fldChar')
                                fldChar.set(qn('w:fldCharType'), 'end')
                                run._element.append(fldChar)
                        elif isinstance(footer_content, str):
                            # Just text
                            section.footer.paragraphs[0].text = footer_content
            
            # Process TOC if we have TOC info
            if hasattr(self, 'toc_info') and self.toc_info:
                # In a proper implementation, we would add a TOC field
                # Python-docx doesn't directly support TOC fields,
                # so we'd need to use the low-level API or a workaround
                # For now, we'll just add a message
                # The first paragraph after the TOC heading is where the TOC would go
                for para in self.document.paragraphs:
                    if para.style.name.startswith('Heading 1') and para.text in [toc['title'] for toc in self.toc_info]:
                        next_para_index = self.document.paragraphs.index(para) + 1
                        if next_para_index < len(self.document.paragraphs):
                            toc_para = self.document.paragraphs[next_para_index]
                            toc_para.text = "[Table of contents will be generated when opened in Word]"
            
            # Save the document with all our changes
            self.document.save(file_path)
            self.current_file = file_path
            self.status_var.set(f"Saved: {os.path.basename(file_path)}")
            messagebox.showinfo("Success", f"Saved to {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save file: {str(e)}")
    
    # Table handling methods
    def insert_table_dialog(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Insert Table")
        dialog.geometry("300x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Row and column selection
        row_frame = ttk.Frame(dialog)
        row_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(row_frame, text="Rows:").pack(side=tk.LEFT, padx=5)
        row_var = tk.StringVar(value="3")
        row_spin = tk.Spinbox(row_frame, from_=1, to=20, textvariable=row_var, width=5)
        row_spin.pack(side=tk.LEFT, padx=5)
        
        col_frame = ttk.Frame(dialog)
        col_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(col_frame, text="Columns:").pack(side=tk.LEFT, padx=5)
        col_var = tk.StringVar(value="3")
        col_spin = tk.Spinbox(col_frame, from_=1, to=10, textvariable=col_var, width=5)
        col_spin.pack(side=tk.LEFT, padx=5)
        
        # Insert function
        def insert_table():
            try:
                rows = int(row_var.get())
                cols = int(col_var.get())
                
                if not self.document:
                    self.document = Document()
                    
                new_table = self.document.add_table(rows=rows, cols=cols)
                new_table.style = 'Table Grid'
                
                if not hasattr(self, 'tables'):
                    self.tables = []
                    
                # Add the table to our tracking list
                self.tables.append(new_table)
                
                # Add a placeholder in the text editor
                table_index = len(self.tables)
                cursor_pos = self.text_editor.index(tk.INSERT)
                self.text_editor.insert(cursor_pos, f"\n[TABLE {table_index}]\n")
                
                self.status_var.set(f"Table inserted with {rows} rows and {cols} columns")
                dialog.destroy()
                
                # Update the document structure view
                self.update_document_structure()
                
                # Show the edit table dialog
                self.current_table = new_table
                self.edit_table_dialog()
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to insert table: {str(e)}")
                dialog.destroy()
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=20)
        
        insert_button = ttk.Button(button_frame, text="Insert", command=insert_table)
        insert_button.pack(side=tk.RIGHT, padx=5)
        
        cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
        cancel_button.pack(side=tk.RIGHT, padx=5)
    
    def edit_table_dialog(self):
        if not hasattr(self, 'tables') or not self.tables:
            messagebox.showinfo("No Tables", "There are no tables in the document to edit.")
            return
        
        # If no current table is selected, ask the user to choose one
        if not self.current_table:
            self.select_table_dialog()
            if not self.current_table:  # User cancelled
                return
        
        dialog = tk.Toplevel(self.root)
        dialog.title("Edit Table")
        dialog.geometry("600x400")
        dialog.transient(self.root)
        
        # Create a simple table editor
        editorFrame = ttk.Frame(dialog)
        editorFrame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create entry widgets for each cell
        for i, row in enumerate(self.current_table.rows):
            for j, cell in enumerate(row.cells):
                entry = ttk.Entry(editorFrame, width=20)
                entry.insert(0, cell.text)
                entry.grid(row=i, column=j, padx=2, pady=2, sticky=tk.NSEW)
        
        # Function to apply changes
        def apply_changes():
            # Get all entries from the grid
            entries = editorFrame.winfo_children()
            for i, row in enumerate(self.current_table.rows):
                for j, cell in enumerate(row.cells):
                    index = i * len(row.cells) + j
                    if index < len(entries):
                        cell.text = entries[index].get()
            
            self.status_var.set("Table updated")
            dialog.destroy()
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        apply_button = ttk.Button(button_frame, text="Apply", command=apply_changes)
        apply_button.pack(side=tk.RIGHT, padx=5)
        
        cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
        cancel_button.pack(side=tk.RIGHT, padx=5)
    
    def select_table_dialog(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Select Table")
        dialog.geometry("300x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        ttk.Label(dialog, text="Select a table to edit:").pack(padx=10, pady=10)
        
        listbox = tk.Listbox(dialog)
        listbox.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        for i, table in enumerate(self.tables):
            listbox.insert(tk.END, f"Table {i+1} ({len(table.rows)}×{len(table.columns)})")
        
        def select_table():
            selected = listbox.curselection()
            if selected:
                index = selected[0]
                self.current_table = self.tables[index]
                dialog.destroy()
        
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        select_button = ttk.Button(button_frame, text="Select", command=select_table)
        select_button.pack(side=tk.RIGHT, padx=5)
        
        cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
        cancel_button.pack(side=tk.RIGHT, padx=5)
    
    def add_table_row(self):
        if not hasattr(self, 'tables') or not self.tables:
            messagebox.showinfo("No Tables", "There are no tables in the document to modify.")
            return
            
        if not self.current_table:
            self.select_table_dialog()
            if not self.current_table:  # User cancelled
                return
        
        # Add a row to the current table
        self.current_table.add_row()
        self.status_var.set("Row added to table")
        self.update_document_structure()
    
    def add_table_column(self):
        if not hasattr(self, 'tables') or not self.tables:
            messagebox.showinfo("No Tables", "There are no tables in the document to modify.")
            return
            
        if not self.current_table:
            self.select_table_dialog()
            if not self.current_table:  # User cancelled
                return
        
        # Get all rows
        rows = self.current_table.rows
        if not rows:
            return
            
        # Add a cell to each row
        for row in rows:
            row.add_cell()
            
        self.status_var.set("Column added to table")
        self.update_document_structure()
    
    def delete_table_row(self):
        if not hasattr(self, 'tables') or not self.tables:
            messagebox.showinfo("No Tables", "There are no tables in the document to modify.")
            return
            
        if not self.current_table:
            self.select_table_dialog()
            if not self.current_table:  # User cancelled
                return
                
        # Let the user select which row to delete
        dialog = tk.Toplevel(self.root)
        dialog.title("Delete Row")
        dialog.geometry("200x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        ttk.Label(dialog, text="Select row to delete:").pack(padx=10, pady=10)
        
        row_var = tk.StringVar(value="1")
        row_spin = tk.Spinbox(dialog, from_=1, to=len(self.current_table.rows), textvariable=row_var, width=5)
        row_spin.pack(padx=10, pady=5)
        
        def delete_row():
            try:
                row_idx = int(row_var.get()) - 1
                if 0 <= row_idx < len(self.current_table.rows):
                    # Delete the row - this is a bit harder as python-docx doesn't have a direct method
                    row = self.current_table.rows[row_idx]._tr
                    row.getparent().remove(row)
                    
                    self.status_var.set(f"Row {row_idx + 1} deleted")
                    self.update_document_structure()
                dialog.destroy()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to delete row: {str(e)}")
                dialog.destroy()
        
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        delete_button = ttk.Button(button_frame, text="Delete", command=delete_row)
        delete_button.pack(side=tk.RIGHT, padx=5)
        
        cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
        cancel_button.pack(side=tk.RIGHT, padx=5)
    
    def delete_table_column(self):
        if not hasattr(self, 'tables') or not self.tables:
            messagebox.showinfo("No Tables", "There are no tables in the document to modify.")
            return
            
        if not self.current_table:
            self.select_table_dialog()
            if not self.current_table:  # User cancelled
                return
                
        # Let the user select which column to delete
        dialog = tk.Toplevel(self.root)
        dialog.title("Delete Column")
        dialog.geometry("200x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        ttk.Label(dialog, text="Select column to delete:").pack(padx=10, pady=10)
        
        # Assume we can get columns count from the first row
        if not self.current_table.rows:
            dialog.destroy()
            return
            
        col_count = len(self.current_table.rows[0].cells)
        
        col_var = tk.StringVar(value="1")
        col_spin = tk.Spinbox(dialog, from_=1, to=col_count, textvariable=col_var, width=5)
        col_spin.pack(padx=10, pady=5)
        
        def delete_column():
            try:
                col_idx = int(col_var.get()) - 1
                if col_idx < 0 or col_idx >= col_count:
                    return
                    
                # Delete the column in each row
                for row in self.current_table.rows:
                    cell = row.cells[col_idx]._tc
                    row._tr.remove(cell)
                    
                self.status_var.set(f"Column {col_idx + 1} deleted")
                self.update_document_structure()
                dialog.destroy()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to delete column: {str(e)}")
                dialog.destroy()
        
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        delete_button = ttk.Button(button_frame, text="Delete", command=delete_column)
        delete_button.pack(side=tk.RIGHT, padx=5)
        
        cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
        cancel_button.pack(side=tk.RIGHT, padx=5)
    
    # Image handling methods
    def insert_image(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Image files", "*.jpg *.jpeg *.png *.gif *.bmp"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                # Create a document if it doesn't exist
                if not self.document:
                    self.document = Document()
                
                # Add the image to the document
                paragraph = self.document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(file_path, width=Inches(4))  # Default width
                
                # Track the image
                if not hasattr(self, 'document_images'):
                    self.document_images = []
                self.document_images.append(file_path)
                
                # Add a placeholder in the text
                img_index = len(self.document_images)
                cursor_pos = self.text_editor.index(tk.INSERT)
                self.text_editor.insert(cursor_pos, f"[IMAGE {img_index}]\n")
                
                self.status_var.set(f"Inserted image: {os.path.basename(file_path)}")
                self.update_document_structure()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to insert image: {str(e)}")
    
    def _get_bold_icon(self):
        # Base64 encoded minimal bold icon
        return b'R0lGODlhEAAQAIABAAAAAP///yH5BAEAAAEALAAAAAAQABAAAAIjjI+py+0Po5wHVIBzVphqa3zbmlFNyYksB4bj+TJoPS0FADs='

    def _get_italic_icon(self):
        # Base64 encoded minimal italic icon
        return b'R0lGODlhEAAQAIABAAAAAP///yH5BAEAAAEALAAAAAAQABAAAAIejI+py+0Po5wGNIBzZZhN24lQBoZkB35jKnIkWwAAOw=='

    def _get_underline_icon(self):
        # Base64 encoded minimal underline icon
        return b'R0lGODlhEAAQAIABAAAAAP///yH5BAEAAAEALAAAAAAQABAAAAIijI+py+0Po5wSgAtzfoCbBXJbGI5MGZ4kdY2qh75zWgAAOw=='

    def _get_align_left_icon(self):
        # Base64 encoded minimal align left icon
        return b'R0lGODlhEAAQAIABAAAAAP///yH5BAEAAAEALAAAAAAQABAAAAIdjI+py+0Po5y02ouz3rz7D4biSJbmiaIqeK5LAQA7'

    def _get_align_center_icon(self):
        # Base64 encoded minimal align center icon
        return b'R0lGODlhEAAQAIABAAAAAP///yH5BAEAAAEALAAAAAAQABAAAAIdjI+py+0Po5y02ouz3jzgD4ZiSJZmiX4qyrbuUwAAOw=='

    def _get_align_right_icon(self):
        # Base64 encoded minimal align right icon
        return b'R0lGODlhEAAQAIABAAAAAP///yH5BAEAAAEALAAAAAAQABAAAAIdjI+py+0Po5y02ouz3rz7D4YiSJbmiZ6purZxVAAAOw=='
    
    # Tooltip class for creating tooltips on hover
    def _create_tooltip(self, widget, text):
        tooltip = ToolTip(widget, text)
        self.tooltips[widget] = tooltip
        return tooltip
    
    # Keyboard shortcuts setup
    def _setup_keyboard_shortcuts(self):
        self.root.bind("<Control-b>", lambda e: self.toggle_bold())
        self.root.bind("<Control-i>", lambda e: self.toggle_italic())
        self.root.bind("<Control-u>", lambda e: self.toggle_underline())
        self.root.bind("<Control-l>", lambda e: self.set_alignment("left"))
        self.root.bind("<Control-e>", lambda e: self.set_alignment("center"))
        self.root.bind("<Control-r>", lambda e: self.set_alignment("right"))
        
        # Standard keyboard shortcuts for file operations
        self.root.bind("<Control-n>", lambda e: self.new_document())
        self.root.bind("<Control-o>", lambda e: self.open_file())
        self.root.bind("<Control-s>", lambda e: self.save_file())
        self.root.bind("<Control-S>", lambda e: self.save_file_as())  # Ctrl+Shift+S
        self.root.bind("<Control-f>", lambda e: self.find_replace_dialog())

    def clear_text(self):
        self.text_editor.delete(1.0, tk.END)

    # Text formatting methods
    def change_font_family(self, event=None):
        selected_font = self.font_family.get()
        self.current_font_family = selected_font
        
        # Apply font to currently selected text if there's a selection
        try:
            selection_start = self.text_editor.index(tk.SEL_FIRST)
            selection_end = self.text_editor.index(tk.SEL_LAST)
            # In a real implementation, would apply font here via tags
        except tk.TclError:
            # No selection, just update state
            pass
            
        self.status_var.set(f"Font changed to {selected_font}")
    
    def change_font_size(self, event=None):
        selected_size = self.font_size.get()
        self.current_font_size = int(selected_size)
        
        # Apply font size to currently selected text if there's a selection
        try:
            selection_start = self.text_editor.index(tk.SEL_FIRST)
            selection_end = self.text_editor.index(tk.SEL_LAST)
            # In a real implementation, would apply font size here via tags
        except tk.TclError:
            # No selection, just update state
            pass
            
        self.status_var.set(f"Font size changed to {selected_size}")
    
    def toggle_bold(self):
        self.current_bold = not self.current_bold
        
        # Update the button state
        self.update_formatting_buttons()
        
        # Apply bold to currently selected text if there's a selection
        try:
            selection_start = self.text_editor.index(tk.SEL_FIRST)
            selection_end = self.text_editor.index(tk.SEL_LAST)
            # In a real implementation, would apply bold here via tags
        except tk.TclError:
            # No selection, just update state
            pass
            
        self.status_var.set(f"Bold formatting {'applied' if self.current_bold else 'removed'}")
    
    def toggle_italic(self):
        self.current_italic = not self.current_italic
        
        # Update the button state
        self.update_formatting_buttons()
        
        # Apply italic to currently selected text if there's a selection
        try:
            selection_start = self.text_editor.index(tk.SEL_FIRST)
            selection_end = self.text_editor.index(tk.SEL_LAST)
            # In a real implementation, would apply italic here via tags
        except tk.TclError:
            # No selection, just update state
            pass
            
        self.status_var.set(f"Italic formatting {'applied' if self.current_italic else 'removed'}")
    
    def toggle_underline(self):
        self.current_underline = not self.current_underline
        
        # Update the button state
        self.update_formatting_buttons()
        
        # Apply underline to currently selected text if there's a selection
        try:
            selection_start = self.text_editor.index(tk.SEL_FIRST)
            selection_end = self.text_editor.index(tk.SEL_LAST)
            # In a real implementation, would apply underline here via tags
        except tk.TclError:
            # No selection, just update state
            pass
            
        self.status_var.set(f"Underline formatting {'applied' if self.current_underline else 'removed'}")
    
    def text_color_dialog(self):
        # Open the color chooser dialog
        color = colorchooser.askcolor(initialcolor=self.current_color, title="Select Text Color")
        
        # If a color was selected (not cancelled)
        if color[1]:
            self.current_color = color[1]
            
            # Update the color button to show the selected color
            try:
                self.color_button.configure(fg=color[1])
            except:
                pass
                
            # Apply color to currently selected text if there's a selection
            try:
                selection_start = self.text_editor.index(tk.SEL_FIRST)
                selection_end = self.text_editor.index(tk.SEL_LAST)
                # In a real implementation, would apply text color here via tags
            except tk.TclError:
                # No selection, just update state
                pass
                
            self.status_var.set(f"Text color changed to {color[1]}")
    
    def set_alignment(self, alignment):
        self.current_alignment = alignment
        
        # Update the button state
        self.update_formatting_buttons()
        
        # Get the current line or paragraph
        try:
            # If there's a selection, align all lines in the selection
            selection_start = self.text_editor.index(tk.SEL_FIRST)
            selection_end = self.text_editor.index(tk.SEL_LAST)
            # Would apply alignment to selected lines here
        except tk.TclError:
            # No selection, get current paragraph
            cursor_pos = self.text_editor.index(tk.INSERT)
            line_start = cursor_pos.split('.')[0] + '.0'
            # Would apply alignment to current paragraph here
            
        self.status_var.set(f"Text alignment set to {alignment}")
    
    # Paragraph handling methods
    def paragraph_style_dialog(self):
        # Create a dialog to select paragraph styles
        dialog = tk.Toplevel(self.root)
        dialog.title("Paragraph Styles")
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        style_listbox = tk.Listbox(dialog)
        style_listbox.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Populate with standard paragraph styles
        styles = ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle", "Quote", "List Paragraph"]
        for style in styles:
            style_listbox.insert(tk.END, style)
        
        # Function to apply the selected style
        def apply_style():
            selected_indices = style_listbox.curselection()
            if selected_indices:
                selected_style = style_listbox.get(selected_indices[0])
                self.status_var.set(f"Applied style: {selected_style}")
                dialog.destroy()
        
        # Apply button
        apply_button = ttk.Button(dialog, text="Apply", command=apply_style)
        apply_button.pack(pady=10)

    # Document structure methods
    def update_document_structure(self):
        # Clear the current structure
        for item in self.structure_tree.get_children():
            self.structure_tree.delete(item)
            
        if not self.document:
            return
            
        # Add document elements to the tree view
        doc_node = self.structure_tree.insert("", "end", text="Document")
        
        # Add sections if we have them tracked
        if hasattr(self, 'sections') and self.sections:
            section_node = self.structure_tree.insert(doc_node, "end", text="Sections")
            for i, section in enumerate(self.document.sections):
                section_item = self.structure_tree.insert(section_node, "end", text=f"Section {i+1}")
                
                # Add header info
                if hasattr(self, 'headers') and i in self.headers:
                    header_text = self.headers[i].text if hasattr(self.headers[i], 'text') else "[Header]" 
                    self.structure_tree.insert(section_item, "end", text=f"Header: {header_text[:30] + '...' if len(header_text) > 30 else header_text}")
                
                # Add footer info
                if hasattr(self, 'footers') and i in self.footers:
                    footer_text = self.footers[i].text if hasattr(self.footers[i], 'text') else "[Footer]"
                    self.structure_tree.insert(section_item, "end", text=f"Footer: {footer_text[:30] + '...' if len(footer_text) > 30 else footer_text}")
        
        # Add paragraph elements
        para_node = self.structure_tree.insert(doc_node, "end", text="Paragraphs")
        for i, para in enumerate(self.document.paragraphs):
            para_text = para.text
            style_name = para.style.name if hasattr(para, 'style') and hasattr(para.style, 'name') else "Normal"
            self.structure_tree.insert(para_node, "end", text=f"Paragraph {i+1} [{style_name}]: {para_text[:30] + '...' if len(para_text) > 30 else para_text}")
        
        # Add table elements if any
        if self.document.tables:
            table_node = self.structure_tree.insert(doc_node, "end", text="Tables")
            for i, table in enumerate(self.document.tables):
                self.structure_tree.insert(table_node, "end", text=f"Table {i+1} ({len(table.rows)}×{len(table.columns)})")
                
        # Add image elements if any
        if hasattr(self, 'document_images') and self.document_images:
            image_node = self.structure_tree.insert(doc_node, "end", text="Images")
            for i, img_path in enumerate(self.document_images):
                self.structure_tree.insert(image_node, "end", text=f"Image {i+1}: {os.path.basename(img_path)}")
        
        # Update the properties tab
        self.update_properties_tab()
    
    def update_properties_tab(self):
        # Clear current properties
        for widget in self.properties_frame.winfo_children():
            widget.destroy()
            
        # Create a frame for properties
        prop_frame = ttk.Frame(self.properties_frame)
        prop_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Document info
        ttk.Label(prop_frame, text="Document Information", font=("Arial", 12, "bold")).grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=5)
        
        # File path
        ttk.Label(prop_frame, text="File Path:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=2)
        ttk.Label(prop_frame, text=self.current_file if self.current_file else "Not saved").grid(row=1, column=1, sticky=tk.W, padx=5, pady=2)
        
        # Document stats
        if self.document:
            # Count paragraphs
            para_count = len(self.document.paragraphs)
            ttk.Label(prop_frame, text="Paragraphs:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=2)
            ttk.Label(prop_frame, text=str(para_count)).grid(row=2, column=1, sticky=tk.W, padx=5, pady=2)
            
            # Count tables
            table_count = len(self.document.tables) if hasattr(self.document, 'tables') else 0
            ttk.Label(prop_frame, text="Tables:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=2)
            ttk.Label(prop_frame, text=str(table_count)).grid(row=3, column=1, sticky=tk.W, padx=5, pady=2)
            
            # Count images
            image_count = len(self.document_images) if hasattr(self, 'document_images') else 0
            ttk.Label(prop_frame, text="Images:").grid(row=4, column=0, sticky=tk.W, padx=5, pady=2)
            ttk.Label(prop_frame, text=str(image_count)).grid(row=4, column=1, sticky=tk.W, padx=5, pady=2)
            
            # Word count (approximate)
            word_count = sum(len(para.text.split()) for para in self.document.paragraphs)
            ttk.Label(prop_frame, text="Word Count (approx):").grid(row=5, column=0, sticky=tk.W, padx=5, pady=2)
            ttk.Label(prop_frame, text=str(word_count)).grid(row=5, column=1, sticky=tk.W, padx=5, pady=2)

    # Font dialog method
    def font_dialog(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Font Selection")
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Font family selection
        font_frame = ttk.LabelFrame(dialog, text="Font")
        font_frame.pack(fill=tk.X, padx=10, pady=5)
        
        font_listbox = tk.Listbox(font_frame, height=5)
        font_listbox.pack(fill=tk.X, padx=10, pady=5)
        
        # Populate with available fonts
        available_fonts = sorted(font.families())
        for f in available_fonts:
            font_listbox.insert(tk.END, f)
            
        # Try to select the current font
        try:
            current_index = available_fonts.index(self.current_font_family)
            font_listbox.selection_set(current_index)
            font_listbox.see(current_index)
        except ValueError:
            pass
        
        # Font size selection
        size_frame = ttk.LabelFrame(dialog, text="Size")
        size_frame.pack(fill=tk.X, padx=10, pady=5)
        
        size_listbox = tk.Listbox(size_frame, height=5)
        size_listbox.pack(fill=tk.X, padx=10, pady=5)
        
        # Populate with common font sizes
        sizes = [8, 9, 10, 11, 12, 14, 16, 18, 20, 22, 24, 26, 28, 36, 48, 72]
        for size in sizes:
            size_listbox.insert(tk.END, size)
            
        # Try to select the current size
        try:
            current_size_index = sizes.index(self.current_font_size)
            size_listbox.selection_set(current_size_index)
            size_listbox.see(current_size_index)
        except ValueError:
            pass
        
        # Style options
        style_frame = ttk.LabelFrame(dialog, text="Style")
        style_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Checkbuttons for bold, italic, underline
        bold_var = tk.BooleanVar(value=self.current_bold)
        bold_check = ttk.Checkbutton(style_frame, text="Bold", variable=bold_var)
        bold_check.pack(side=tk.LEFT, padx=10, pady=5)
        
        italic_var = tk.BooleanVar(value=self.current_italic)
        italic_check = ttk.Checkbutton(style_frame, text="Italic", variable=italic_var)
        italic_check.pack(side=tk.LEFT, padx=10, pady=5)
        
        underline_var = tk.BooleanVar(value=self.current_underline)
        underline_check = ttk.Checkbutton(style_frame, text="Underline", variable=underline_var)
        underline_check.pack(side=tk.LEFT, padx=10, pady=5)
        
        # Apply function
        def apply_font():
            # Get font family
            font_selection = font_listbox.curselection()
            if font_selection:
                self.current_font_family = available_fonts[font_selection[0]]
            
            # Get font size
            size_selection = size_listbox.curselection()
            if size_selection:
                self.current_font_size = sizes[size_selection[0]]
            
            # Get styles
            self.current_bold = bold_var.get()
            self.current_italic = italic_var.get()
            self.current_underline = underline_var.get()
            
            self.status_var.set(f"Font updated: {self.current_font_family}, {self.current_font_size}pt")
            dialog.destroy()
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        apply_button = ttk.Button(button_frame, text="Apply", command=apply_font)
        apply_button.pack(side=tk.RIGHT, padx=5)
        
        cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
        cancel_button.pack(side=tk.RIGHT, padx=5)
    
    # Additional functionality methods
    def undo(self):
        try:
            self.text_editor.edit_undo()
            self.status_var.set("Undo successful")
        except tk.TclError:
            self.status_var.set("Nothing to undo")
    
    def redo(self):
        try:
            self.text_editor.edit_redo()
            self.status_var.set("Redo successful")
        except tk.TclError:
            self.status_var.set("Nothing to redo")
    
    def find_replace_dialog(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Find and Replace")
        dialog.geometry("400x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Find frame
        find_frame = ttk.LabelFrame(dialog, text="Find")
        find_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(find_frame, text="Find what:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        find_entry = ttk.Entry(find_frame, width=30)
        find_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        find_entry.focus_set()
        
        # Replace frame
        replace_frame = ttk.LabelFrame(dialog, text="Replace")
        replace_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(replace_frame, text="Replace with:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        replace_entry = ttk.Entry(replace_frame, width=30)
        replace_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        
        # Options frame
        options_frame = ttk.Frame(dialog)
        options_frame.pack(fill=tk.X, padx=10, pady=5)
        
        case_sensitive = tk.BooleanVar(value=False)
        case_check = ttk.Checkbutton(options_frame, text="Case sensitive", variable=case_sensitive)
        case_check.pack(side=tk.LEFT, padx=5)
        
        whole_word = tk.BooleanVar(value=False)
        word_check = ttk.Checkbutton(options_frame, text="Whole word", variable=whole_word)
        word_check.pack(side=tk.LEFT, padx=5)
        
        # Button functions
        def find_text():
            # Clear any existing tags
            self.text_editor.tag_remove('found', '1.0', tk.END)
            
            search_text = find_entry.get()
            if not search_text:
                return
                
            start_pos = '1.0'
            count_var = tk.StringVar()
            
            while True:
                # Find the text, considering options
                if case_sensitive.get():
                    start_pos = self.text_editor.search(search_text, start_pos, stopindex=tk.END, count=count_var, nocase=0)
                else:
                    start_pos = self.text_editor.search(search_text, start_pos, stopindex=tk.END, count=count_var, nocase=1)
                    
                if not start_pos:
                    break
                    
                # Calculate end position
                end_pos = f"{start_pos}+{count_var.get()}c"
                
                # Check if it's a whole word match if option is selected
                if whole_word.get():
                    # Get characters before and after the match
                    # Don't do boundary check for start of text
                    if start_pos != "1.0":
                        before_pos = f"{start_pos}-1c"
                        before_char = self.text_editor.get(before_pos, start_pos)
                        if before_char.isalnum() or before_char == '_':
                            start_pos = end_pos  # Not a word boundary, continue search
                            continue
                            
                    # Don't do boundary check for end of text
                    if end_pos != tk.END:
                        after_char = self.text_editor.get(end_pos, f"{end_pos}+1c")
                        if after_char.isalnum() or after_char == '_':
                            start_pos = end_pos  # Not a word boundary, continue search
                            continue
                
                # Add tag to highlight the match
                self.text_editor.tag_add('found', start_pos, end_pos)
                start_pos = end_pos  # Move to the end of the current match
            
            # Configure the tag for highlighting
            self.text_editor.tag_config('found', background='yellow', foreground='black')
            
            # Check if any matches were found
            if self.text_editor.tag_ranges('found'):
                # Scroll to the first match
                self.text_editor.see(self.text_editor.tag_ranges('found')[0])
                self.status_var.set(f"Found matches for '{search_text}'")
            else:
                self.status_var.set(f"No matches found for '{search_text}'")
        
        def replace_text():
            # Get current selection or find next match if none selected
            try:
                selected_text = self.text_editor.get(tk.SEL_FIRST, tk.SEL_LAST)
                if selected_text == find_entry.get():
                    self.text_editor.delete(tk.SEL_FIRST, tk.SEL_LAST)
                    self.text_editor.insert(tk.SEL_FIRST, replace_entry.get())
                    self.status_var.set("Replaced selection")
            except tk.TclError:
                # No selection, find and replace next match
                find_text()  # Find and highlight next match
                
                # Check if any matches were found
                if self.text_editor.tag_ranges('found'):
                    # Get the position of the first match
                    start_pos = self.text_editor.tag_ranges('found')[0]
                    end_pos = self.text_editor.tag_ranges('found')[1]
                    
                    # Replace the text
                    self.text_editor.delete(start_pos, end_pos)
                    self.text_editor.insert(start_pos, replace_entry.get())
                    self.status_var.set("Replaced match")
        
        def replace_all():
            # Clear any existing tags
            self.text_editor.tag_remove('found', '1.0', tk.END)
            
            search_text = find_entry.get()
            replace_text = replace_entry.get()
            if not search_text:
                return
                
            # Keep track of replacements
            count = 0
            start_pos = '1.0'
            count_var = tk.StringVar()
            
            # Disable redrawing to improve performance during multiple replacements
            self.text_editor.config(state=tk.NORMAL)
            
            while True:
                # Find the text, considering options
                if case_sensitive.get():
                    start_pos = self.text_editor.search(search_text, start_pos, stopindex=tk.END, count=count_var, nocase=0)
                else:
                    start_pos = self.text_editor.search(search_text, start_pos, stopindex=tk.END, count=count_var, nocase=1)
                    
                if not start_pos:
                    break
                    
                # Calculate end position
                end_pos = f"{start_pos}+{count_var.get()}c"
                
                # Check if it's a whole word match if option is selected
                if whole_word.get():
                    # Get characters before and after the match
                    before_match = False
                    after_match = False
                    
                    # Check before match
                    if start_pos != "1.0":
                        before_pos = f"{start_pos}-1c"
                        before_char = self.text_editor.get(before_pos, start_pos)
                        before_match = before_char.isalnum() or before_char == '_'
                        
                    # Check after match
                    if end_pos != tk.END:
                        after_char = self.text_editor.get(end_pos, f"{end_pos}+1c")
                        after_match = after_char.isalnum() or after_char == '_'
                        
                    if before_match or after_match:
                        start_pos = end_pos  # Not a word boundary, continue search
                        continue
                
                # Replace the text
                self.text_editor.delete(start_pos, end_pos)
                self.text_editor.insert(start_pos, replace_text)
                count += 1
                
                # Update the start position for the next search
                # Need to adjust it based on the replacement length
                start_pos = f"{start_pos}+{len(replace_text)}c"
            
            self.status_var.set(f"Replaced {count} occurrences")
        
        # Button frame
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        find_button = ttk.Button(button_frame, text="Find Next", command=find_text)
        find_button.pack(side=tk.LEFT, padx=5)
        
        replace_button = ttk.Button(button_frame, text="Replace", command=replace_text)
        replace_button.pack(side=tk.LEFT, padx=5)
        
        replace_all_button = ttk.Button(button_frame, text="Replace All", command=replace_all)
        replace_all_button.pack(side=tk.LEFT, padx=5)
        
        close_button = ttk.Button(button_frame, text="Close", command=dialog.destroy)
        close_button.pack(side=tk.RIGHT, padx=5)
    
    def insert_page_break(self):
        if not self.document:
            self.document = Document()
            
        cursor_pos = self.text_editor.index(tk.INSERT)
        self.text_editor.insert(cursor_pos, "\f")
        self.status_var.set("Page break inserted")
    
    def insert_section_break(self):
        if not self.document:
            self.document = Document()
            
        cursor_pos = self.text_editor.index(tk.INSERT)
        self.text_editor.insert(cursor_pos, "\n[SECTION BREAK]\n")
        
        # Track section for when we save
        if not hasattr(self, 'section_breaks'):
            self.section_breaks = []
            
        # Save the position where we inserted the section break
        line_number = int(cursor_pos.split('.')[0])
        self.section_breaks.append(line_number)
        
        self.status_var.set("Section break inserted")
    
    def insert_hyperlink(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Insert Hyperlink")
        dialog.geometry("400x150")
        dialog.transient(self.root)
        dialog.grab_set()
        
        ttk.Label(dialog, text="Text to display:").grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        text_entry = ttk.Entry(dialog, width=30)
        text_entry.grid(row=0, column=1, padx=10, pady=10, sticky=tk.W)
        text_entry.focus_set()
        
        ttk.Label(dialog, text="URL:").grid(row=1, column=0, padx=10, pady=10, sticky=tk.W)
        url_entry = ttk.Entry(dialog, width=30)
        url_entry.grid(row=1, column=1, padx=10, pady=10, sticky=tk.W)
        
        def insert_link():
            text = text_entry.get()
            url = url_entry.get()
            
            if not text or not url:
                messagebox.showwarning("Missing Information", "Both text and URL are required")
                return
                
            # Insert a placeholder in the text editor
            cursor_pos = self.text_editor.index(tk.INSERT)
            self.text_editor.insert(cursor_pos, f"[LINK: {text}]")
            
            # Store the link information for when saving the document
            if not hasattr(self, 'hyperlinks'):
                self.hyperlinks = []
                
            self.hyperlinks.append((text, url))
            self.status_var.set(f"Hyperlink to {url} inserted")
            dialog.destroy()
        
        button_frame = ttk.Frame(dialog)
        button_frame.grid(row=2, column=0, columnspan=2, pady=10)
        
        insert_button = ttk.Button(button_frame, text="Insert", command=insert_link)
        insert_button.pack(side=tk.LEFT, padx=10)
        
        cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
        cancel_button.pack(side=tk.LEFT, padx=10)
    
    def export_pdf(self):
        if not self.document:
            messagebox.showinfo("No Document", "Please open or create a document first.")
            return
            
        if not self.current_file:
            messagebox.showinfo("Save Required", "Please save the document first before exporting to PDF.")
            self.save_file_as()
            if not self.current_file:  # User cancelled
                return
                
        # Get the output PDF file path
        pdf_file = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
            initialfile=os.path.splitext(os.path.basename(self.current_file))[0] + ".pdf"
        )
        
        if not pdf_file:
            return
            
        try:
            # We can't directly convert to PDF with python-docx, so inform the user
            messagebox.showinfo("PDF Export", "Direct PDF export is not supported in this version.\n\n"
                              "Your document has been saved, and can be opened in Microsoft Word or LibreOffice "
                              "where you can export it as PDF.")
            self.status_var.set("Document saved and ready for PDF export")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to prepare for PDF export: {str(e)}")
    
    def document_properties(self):
        # Show document properties tab
        self.notebook.select(2)  # Index 2 should be the Properties tab
    
    def zoom_in(self):
        # Change font size for the text editor
        current_font = font.Font(font=self.text_editor["font"])
        size = current_font.actual()["size"] + 1
        self.text_editor.config(font=(current_font.actual()["family"], size))
        self.status_var.set(f"Zoom level: {size}")
    
    def zoom_out(self):
        # Change font size for the text editor
        current_font = font.Font(font=self.text_editor["font"])
        size = max(8, current_font.actual()["size"] - 1)  # Don't go below 8pt
        self.text_editor.config(font=(current_font.actual()["family"], size))
        self.status_var.set(f"Zoom level: {size}")
    
    def reset_zoom(self):
        # Reset to default font size
        current_font = font.Font(font=self.text_editor["font"])
        self.text_editor.config(font=(current_font.actual()["family"], 10))
        self.status_var.set("Zoom level: 10 (default)")
    
    def show_about(self):
        messagebox.showinfo("About DOCX Editor", 
                          "DOCX Editor\n\n"
                          "Version 1.0\n\n"
                          "A Python application for editing Microsoft Word documents.\n\n"
                          "Features:\n"
                          "- Text formatting\n"
                          "- Table support\n"
                          "- Image handling\n"
                          "- Paragraph styling\n\n"
                          "Built with python-docx")
    
    def edit_header(self):
        if not self.document:
            messagebox.showinfo("No Document", "Please open or create a document first.")
            return
            
        # Create a dialog for editing the header
        dialog = tk.Toplevel(self.root)
        dialog.title("Edit Header")
        dialog.geometry("600x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Section selection if document has multiple sections
        section_frame = ttk.Frame(dialog)
        section_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(section_frame, text="Section:").pack(side=tk.LEFT, padx=5)
        section_var = tk.StringVar(value="1")
        
        # Get section count
        section_count = len(self.document.sections) if hasattr(self.document, 'sections') else 1
        section_spin = tk.Spinbox(section_frame, from_=1, to=section_count, textvariable=section_var, width=5)
        section_spin.pack(side=tk.LEFT, padx=5)
        
        # Editor for header content
        editor_frame = ttk.LabelFrame(dialog, text="Header Content")
        editor_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        header_editor = scrolledtext.ScrolledText(editor_frame, wrap=tk.WORD, height=8)
        header_editor.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Try to get existing header for this section
        def get_header_content():
            try:
                section_idx = int(section_var.get()) - 1
                if 0 <= section_idx < section_count:
                    section = self.document.sections[section_idx]
                    header = section.header
                    header_text = ""
                    
                    # Extract header paragraphs
                    for para in header.paragraphs:
                        header_text += para.text + "\n"
                    
                    # Update editor
                    header_editor.delete(1.0, tk.END)
                    header_editor.insert(tk.END, header_text)
                    
                    # Store the header
                    self.headers[section_idx] = header
            except Exception as e:
                messagebox.showinfo("Header Info", "This section doesn't have a header yet. You can add content now.")
        
        # Initial load
        get_header_content()
        
        # If section is changed, load the corresponding header
        section_spin.configure(command=get_header_content)
        
        # Function to apply header changes
        def apply_header():
            try:
                section_idx = int(section_var.get()) - 1
                if 0 <= section_idx < section_count:
                    # Get the content
                    header_text = header_editor.get(1.0, tk.END)
                    
                    # We can't directly modify the header here, but we'll store the text
                    # and apply it when saving the document
                    self.headers[section_idx] = header_text
                    
                    self.status_var.set(f"Header updated for section {section_idx + 1}")
                    
                    # Update the document structure view
                    self.update_document_structure()
                    
                    dialog.destroy()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to update header: {str(e)}")
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        apply_button = ttk.Button(button_frame, text="Apply", command=apply_header)
        apply_button.pack(side=tk.RIGHT, padx=5)
        
        cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
        cancel_button.pack(side=tk.RIGHT, padx=5)
    
    def edit_footer(self):
        if not self.document:
            messagebox.showinfo("No Document", "Please open or create a document first.")
            return
            
        # Create a dialog for editing the footer
        dialog = tk.Toplevel(self.root)
        dialog.title("Edit Footer")
        dialog.geometry("600x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Section selection if document has multiple sections
        section_frame = ttk.Frame(dialog)
        section_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(section_frame, text="Section:").pack(side=tk.LEFT, padx=5)
        section_var = tk.StringVar(value="1")
        
        # Get section count
        section_count = len(self.document.sections) if hasattr(self.document, 'sections') else 1
        section_spin = tk.Spinbox(section_frame, from_=1, to=section_count, textvariable=section_var, width=5)
        section_spin.pack(side=tk.LEFT, padx=5)
        
        # Add checkbox for page numbers
        page_num_var = tk.BooleanVar(value=False)
        page_num_check = ttk.Checkbutton(section_frame, text="Include page numbers", variable=page_num_var)
        page_num_check.pack(side=tk.LEFT, padx=20)
        
        # Editor for footer content
        editor_frame = ttk.LabelFrame(dialog, text="Footer Content")
        editor_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        footer_editor = scrolledtext.ScrolledText(editor_frame, wrap=tk.WORD, height=8)
        footer_editor.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Try to get existing footer for this section
        def get_footer_content():
            try:
                section_idx = int(section_var.get()) - 1
                if 0 <= section_idx < section_count:
                    section = self.document.sections[section_idx]
                    footer = section.footer
                    footer_text = ""
                    
                    # Extract footer paragraphs
                    for para in footer.paragraphs:
                        footer_text += para.text + "\n"
                    
                    # Update editor
                    footer_editor.delete(1.0, tk.END)
                    footer_editor.insert(tk.END, footer_text)
                    
                    # Store the footer
                    self.footers[section_idx] = footer
            except Exception as e:
                messagebox.showinfo("Footer Info", "This section doesn't have a footer yet. You can add content now.")
        
        # Initial load
        get_footer_content()
        
        # If section is changed, load the corresponding footer
        section_spin.configure(command=get_footer_content)
        
        # Function to apply footer changes
        def apply_footer():
            try:
                section_idx = int(section_var.get()) - 1
                if 0 <= section_idx < section_count:
                    # Get the content
                    footer_text = footer_editor.get(1.0, tk.END)
                    
                    # Store information about page numbers
                    footer_data = {
                        'text': footer_text,
                        'page_numbers': page_num_var.get()
                    }
                    
                    # We can't directly modify the footer here, but we'll store the text
                    # and apply it when saving the document
                    self.footers[section_idx] = footer_data
                    
                    self.status_var.set(f"Footer updated for section {section_idx + 1}")
                    
                    # Update the document structure view
                    self.update_document_structure()
                    
                    dialog.destroy()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to update footer: {str(e)}")
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        apply_button = ttk.Button(button_frame, text="Apply", command=apply_footer)
        apply_button.pack(side=tk.RIGHT, padx=5)
        
        cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
        cancel_button.pack(side=tk.RIGHT, padx=5)
    
    def insert_toc(self):
        """Insert a table of contents"""
        if not self.document:
            messagebox.showinfo("No Document", "Please open or create a document first.")
            return
        
        dialog = tk.Toplevel(self.root)
        dialog.title("Insert Table of Contents")
        dialog.geometry("400x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Options frame
        options_frame = ttk.LabelFrame(dialog, text="TOC Options")
        options_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Title for TOC
        title_frame = ttk.Frame(options_frame)
        title_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(title_frame, text="Title:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        title_var = tk.StringVar(value="Table of Contents")
        title_entry = ttk.Entry(title_frame, textvariable=title_var, width=30)
        title_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        
        # Depth of headings to include
        depth_frame = ttk.Frame(options_frame)
        depth_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(depth_frame, text="Depth:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        depth_var = tk.StringVar(value="3")
        depth_spin = tk.Spinbox(depth_frame, from_=1, to=6, textvariable=depth_var, width=5)
        depth_spin.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        
        # Function to insert TOC
        def insert_toc_content():
            try:
                title = title_var.get()
                depth = int(depth_var.get())
                
                # Insert a placeholder in the document
                cursor_pos = self.text_editor.index(tk.INSERT)
                toc_placeholder = f"\n[TABLE OF CONTENTS: {title}, Depth={depth}]\n"
                self.text_editor.insert(cursor_pos, toc_placeholder)
                
                # Store TOC information for when saving
                if not hasattr(self, 'toc_info'):
                    self.toc_info = []
                    
                self.toc_info.append({
                    'title': title,
                    'depth': depth,
                    'position': cursor_pos
                })
                
                self.status_var.set("Table of contents inserted")
                dialog.destroy()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to insert table of contents: {str(e)}")
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        insert_button = ttk.Button(button_frame, text="Insert", command=insert_toc_content)
        insert_button.pack(side=tk.RIGHT, padx=5)
        
        cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
        cancel_button.pack(side=tk.RIGHT, padx=5)
    
    # Method to update button states based on current formatting
    def update_formatting_buttons(self):
        # Highlight the bold button if bold is active
        if self.current_bold:
            self.bold_button.state(['pressed'])
        else:
            self.bold_button.state(['!pressed'])
            
        # Update italic button
        if self.current_italic:
            self.italic_button.state(['pressed'])
        else:
            self.italic_button.state(['!pressed'])
            
        # Update underline button
        if self.current_underline:
            self.underline_button.state(['pressed'])
        else:
            self.underline_button.state(['!pressed'])
            
        # Update alignment buttons
        self.align_left_button.state(['!pressed'])
        self.align_center_button.state(['!pressed'])
        self.align_right_button.state(['!pressed'])
        
        if self.current_alignment == "left":
            self.align_left_button.state(['pressed'])
        elif self.current_alignment == "center":
            self.align_center_button.state(['pressed'])
        elif self.current_alignment == "right":
            self.align_right_button.state(['pressed'])
    
    # Apply paragraph style from dropdown
    def apply_style(self, event=None):
        selected_style = self.style_combo.get()
        cursor_pos = self.text_editor.index(tk.INSERT)
        line_start = cursor_pos.split('.')[0] + '.0'
        line_end = cursor_pos.split('.')[0] + '.end'
        
        # Get the current line
        current_line = self.text_editor.get(line_start, line_end)
        
        # Remove any existing markdown-style heading markers
        if current_line.startswith('# ') or current_line.startswith('## ') or current_line.startswith('### '):
            # Extract the text after the heading marker
            if current_line.startswith('### '):
                text = current_line[4:]
            elif current_line.startswith('## '):
                text = current_line[3:]
            elif current_line.startswith('# '):
                text = current_line[2:]
            
            # Replace the line with clean text
            self.text_editor.delete(line_start, line_end)
            self.text_editor.insert(line_start, text)
        
        # Now apply the new style using markdown-style markers
        if selected_style == "Heading 1":
            self.text_editor.delete(line_start, line_end)
            self.text_editor.insert(line_start, f"# {current_line}")
        elif selected_style == "Heading 2":
            self.text_editor.delete(line_start, line_end)
            self.text_editor.insert(line_start, f"## {current_line}")
        elif selected_style == "Heading 3":
            self.text_editor.delete(line_start, line_end)
            self.text_editor.insert(line_start, f"### {current_line}")
        
        self.status_var.set(f"Applied style: {selected_style}")
        
        # Update the headings navigation dropdown
        self.update_headings_navigation()
    
    # Update list of headings for the navigation dropdown
    def update_headings_navigation(self):
        # Clear the current list
        self.heading_nav.set('')
        headings = []
        
        # Get all lines from the document
        all_text = self.text_editor.get(1.0, tk.END)
        lines = all_text.split('\n')
        
        # Extract headings and their line numbers
        for i, line in enumerate(lines):
            line_number = i + 1  # 1-based line numbering
            if line.startswith('# '):
                headings.append((f"H1: {line[2:]}", line_number))
            elif line.startswith('## '):
                headings.append((f"H2: {line[3:]}", line_number))
            elif line.startswith('### '):
                headings.append((f"H3: {line[4:]}", line_number))
        
        # Update the dropdown values
        self.heading_nav['values'] = [h[0] for h in headings]
        
        # Store the line numbers for navigation
        self.heading_line_numbers = {h[0]: h[1] for h in headings}
    
    # Navigate to a heading when selected from dropdown
    def navigate_to_heading(self, event=None):
        selected = self.heading_nav.get()
        if selected and hasattr(self, 'heading_line_numbers') and selected in self.heading_line_numbers:
            line_number = self.heading_line_numbers[selected]
            self.text_editor.see(f"{line_number}.0")
            self.text_editor.mark_set(tk.INSERT, f"{line_number}.0")
            self.text_editor.focus_set()
            self.status_var.set(f"Navigated to: {selected}")
    
    # Handle highlight color selection
    def highlight_color_dialog(self):
        color = colorchooser.askcolor(initialcolor=self.current_highlight_color, title="Select Highlight Color")
        if color[1]:
            self.current_highlight_color = color[1]
            
            # Update the highlight button color
            try:
                self.highlight_button.configure(bg=color[1])
            except:
                pass
                
            self.status_var.set(f"Highlight color changed to {color[1]}")
    
    # Insert a bullet list at cursor position
    def insert_bullet_list(self):
        cursor_pos = self.text_editor.index(tk.INSERT)
        self.text_editor.insert(cursor_pos, "\n\u2022 Item 1\n\u2022 Item 2\n\u2022 Item 3\n")
        self.status_var.set("Bullet list inserted")
    
    # Insert a numbered list at cursor position
    def insert_numbered_list(self):
        cursor_pos = self.text_editor.index(tk.INSERT)
        self.text_editor.insert(cursor_pos, "\n1. Item 1\n2. Item 2\n3. Item 3\n")
        self.status_var.set("Numbered list inserted")

    def show_help(self):
        help_text = """DOCX Editor Help

Basic Usage:
- Use File menu to create, open, and save documents
- Use formatting toolbar to apply text formatting
- Use Insert menu to add tables, images, and other elements
- Use Table menu to modify tables

Advanced Features:
- Headers and Footers: Use Insert menu to add and edit
- Section Breaks: Divide your document into sections with different layouts
- Table of Contents: Automatically generate based on heading styles
- Page Numbers: Add page numbers to footers

Shortcuts:
- Ctrl+N: New document
- Ctrl+O: Open document
- Ctrl+S: Save document
- Ctrl+Z: Undo
- Ctrl+Y: Redo
- Ctrl+F: Find/Replace

For more help, refer to the documentation.
"""
        
        help_dialog = tk.Toplevel(self.root)
        help_dialog.title("Help")
        help_dialog.geometry("500x400")
        help_dialog.transient(self.root)
        
        help_text_widget = scrolledtext.ScrolledText(help_dialog, wrap=tk.WORD)
        help_text_widget.pack(expand=True, fill=tk.BOTH, padx=10, pady=10)
        help_text_widget.insert(tk.END, help_text)
        help_text_widget.config(state=tk.DISABLED)  # Make it read-only

def main():
    root = tk.Tk()
    app = DocxEditor(root)
    root.mainloop()

if __name__ == "__main__":
    main()
